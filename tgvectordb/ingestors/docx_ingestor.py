from pathlib import Path


def can_handle(filepath: str) -> bool:
    ext = Path(filepath).suffix.lower()
    return ext in (".docx",)


def extract_text(filepath: str, include_headers_footers: bool = False) -> str:
    path = Path(filepath)
    if not path.exists():
        raise FileNotFoundError(f"file not found: {filepath}")
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is needed for .docx support. install with:\n"
            "  pip install python-docx\n"
            "  or: pip install tgvectordb[docx]"
        )
    doc = Document(str(path))
    text_parts = []
    for element in doc.element.body:
        tag = element.tag.split("}")[-1]
        if tag == "p":
            para_text = _extract_paragraph_text(element, doc)
            if para_text:
                text_parts.append(para_text)
        elif tag == "tbl":
            table_text = _extract_table_text(element, doc)
            if table_text:
                text_parts.append(table_text)
    if include_headers_footers:
        hf_text = _extract_headers_footers(doc)
        if hf_text:
            text_parts.insert(0, hf_text)
    return "\n\n".join(text_parts)


def _extract_paragraph_text(element, doc) -> str:
    from docx.oxml.ns import qn

    pPr = element.find(qn("w:pPr"))
    style_name = ""
    if pPr is not None:
        pStyle = pPr.find(qn("w:pStyle"))
        if pStyle is not None:
            style_name = pStyle.get(qn("w:val"), "")
    runs_text = []
    for run in element.iter(qn("w:t")):
        if run.text:
            runs_text.append(run.text)
    full_text = "".join(runs_text).strip()
    if not full_text:
        return ""
    is_list_item = False
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            is_list_item = True
    if style_name.lower().startswith("heading"):
        level = "".join(c for c in style_name if c.isdigit())
        if level:
            prefix = "#" * int(level)
            return f"{prefix} {full_text}"
        else:
            return f"# {full_text}"
    elif style_name.lower() in ("title",):
        return f"# {full_text}"
    elif style_name.lower() in ("subtitle",):
        return f"## {full_text}"
    elif is_list_item:
        return f"- {full_text}"
    else:
        return full_text


def _extract_table_text(element, doc) -> str:
    from docx.oxml.ns import qn

    rows_data = []
    for tr in element.iter(qn("w:tr")):
        cells = []
        for tc in tr.iter(qn("w:tc")):
            cell_texts = []
            for para in tc.iter(qn("w:p")):
                runs = []
                for run in para.iter(qn("w:t")):
                    if run.text:
                        runs.append(run.text)
                cell_text = "".join(runs).strip()
                if cell_text:
                    cell_texts.append(cell_text)
            cells.append(" ".join(cell_texts))
        rows_data.append(cells)
    if not rows_data:
        return ""
    headers = rows_data[0]
    looks_like_headers = all(
        len(h) < 50 and not h.replace(".", "").replace(",", "").isdigit()
        for h in headers
        if h
    )
    result_rows = []
    if looks_like_headers and len(rows_data) > 1:
        for row in rows_data[1:]:
            parts = []
            for i, cell in enumerate(row):
                if not cell:
                    continue
                if i < len(headers) and headers[i]:
                    parts.append(f"{headers[i]}: {cell}")
                else:
                    parts.append(cell)
            if parts:
                result_rows.append(". ".join(parts))
    else:
        for row in rows_data:
            row_text = " | ".join(cell for cell in row if cell)
            if row_text:
                result_rows.append(row_text)
    if result_rows:
        return "Table:\n" + "\n".join(result_rows)
    return ""


def _extract_headers_footers(doc) -> str:
    hf_parts = []
    for section in doc.sections:
        if section.header and not section.header.is_linked_to_previous:
            for para in section.header.paragraphs:
                if para.text.strip():
                    hf_parts.append(para.text.strip())
        if section.footer and not section.footer.is_linked_to_previous:
            for para in section.footer.paragraphs:
                if para.text.strip():
                    hf_parts.append(para.text.strip())
    if hf_parts:
        return "Document header/footer: " + " | ".join(hf_parts)
    return ""
