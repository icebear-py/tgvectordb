"""
tests for source ingestors.
creates temp files and checks that text extraction works correctly.

    pytest tests/test_ingestors.py -v
"""

import os
import json
import tempfile
from pathlib import Path

from tgvectordb.ingestors import text_ingestor, docx_ingestor, pdf_ingestor
from tgvectordb.ingestors.registry import ingest, is_supported, extract_raw_text


class TestTextIngestor:
    def test_plain_txt(self):
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
            f.write("Hello world this is a test document.\n\nSecond paragraph here.")
            f.flush()
            text = text_ingestor.extract_text(f.name)
            assert "Hello world" in text
            assert "Second paragraph" in text
            os.unlink(f.name)

    def test_csv_to_readable(self):
        csv_content = "Name,Age,City\nAlice,25,Mumbai\nBob,30,Delhi\n"
        with tempfile.NamedTemporaryFile(mode="w", suffix=".csv", delete=False) as f:
            f.write(csv_content)
            f.flush()
            text = text_ingestor.extract_text(f.name)
            # should turn into readable format, not raw csv
            assert "Name: Alice" in text
            assert "Age: 25" in text
            os.unlink(f.name)

    def test_html_strips_tags(self):
        html = "<html><body><h1>Title</h1><p>Some <b>bold</b> text.</p><script>evil()</script></body></html>"
        with tempfile.NamedTemporaryFile(mode="w", suffix=".html", delete=False) as f:
            f.write(html)
            f.flush()
            text = text_ingestor.extract_text(f.name)
            assert "Title" in text
            assert "bold" in text
            assert "<h1>" not in text
            assert "evil()" not in text  # script should be removed
            os.unlink(f.name)

    def test_json_text_extraction(self):
        data = [
            {"title": "First Post", "content": "This is the body of the first post with enough text to find."},
            {"title": "Second Post", "content": "Another post with different content and topics."},
        ]
        with tempfile.NamedTemporaryFile(mode="w", suffix=".json", delete=False) as f:
            json.dump(data, f)
            f.flush()
            text = text_ingestor.extract_text(f.name)
            assert "first post" in text.lower()
            assert "second post" in text.lower()
            os.unlink(f.name)

    def test_jsonl_extraction(self):
        lines = [
            json.dumps({"text": "First line of data"}),
            json.dumps({"text": "Second line of data"}),
        ]
        with tempfile.NamedTemporaryFile(mode="w", suffix=".jsonl", delete=False) as f:
            f.write("\n".join(lines))
            f.flush()
            text = text_ingestor.extract_text(f.name)
            assert "First line" in text
            assert "Second line" in text
            os.unlink(f.name)

    def test_code_file_adds_filename(self):
        code = "def hello():\n    print('hello world')\n"
        with tempfile.NamedTemporaryFile(mode="w", suffix=".py", delete=False) as f:
            f.write(code)
            f.flush()
            text = text_ingestor.extract_text(f.name)
            # should include the filename for context
            assert "File:" in text
            assert "hello" in text
            os.unlink(f.name)

    def test_markdown(self):
        md = "# My Document\n\nThis is a paragraph.\n\n## Section Two\n\nMore content here."
        with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False) as f:
            f.write(md)
            f.flush()
            text = text_ingestor.extract_text(f.name)
            assert "My Document" in text
            assert "Section Two" in text
            os.unlink(f.name)

    def test_can_handle_checks(self):
        assert text_ingestor.can_handle("file.txt") == True
        assert text_ingestor.can_handle("file.md") == True
        assert text_ingestor.can_handle("file.csv") == True
        assert text_ingestor.can_handle("file.html") == True
        assert text_ingestor.can_handle("file.py") == True
        assert text_ingestor.can_handle("file.pdf") == False  # pdf has its own ingestor
        assert text_ingestor.can_handle("file.docx") == False  # docx has its own


class TestDocxIngestor:
    def test_can_handle(self):
        assert docx_ingestor.can_handle("report.docx") == True
        assert docx_ingestor.can_handle("report.doc") == False  # old format not supported
        assert docx_ingestor.can_handle("report.pdf") == False

    def test_extract_from_real_docx(self):
        """only runs if python-docx is installed."""
        try:
            from docx import Document
        except ImportError:
            return  # skip if not installed

        # create a simple docx file
        doc = Document()
        doc.add_heading("Test Document", level=1)
        doc.add_paragraph("This is the first paragraph of our test document.")
        doc.add_heading("Section Two", level=2)
        doc.add_paragraph("Second section has different content about machine learning.")

        # add a simple table
        table = doc.add_table(rows=3, cols=2)
        table.cell(0, 0).text = "Name"
        table.cell(0, 1).text = "Score"
        table.cell(1, 0).text = "Alice"
        table.cell(1, 1).text = "95"
        table.cell(2, 0).text = "Bob"
        table.cell(2, 1).text = "87"

        doc.add_paragraph("Final paragraph wrapping things up.")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            text = docx_ingestor.extract_text(f.name)

            assert "Test Document" in text
            assert "first paragraph" in text
            assert "Section Two" in text
            assert "machine learning" in text
            # table content should be there
            assert "Alice" in text
            assert "95" in text or "Score" in text
            assert "Final paragraph" in text

            os.unlink(f.name)

    def test_docx_with_lists(self):
        """test that list items get extracted."""
        try:
            from docx import Document
        except ImportError:
            return

        doc = Document()
        doc.add_paragraph("Shopping list:")
        doc.add_paragraph("Apples", style="List Bullet")
        doc.add_paragraph("Bananas", style="List Bullet")
        doc.add_paragraph("Milk", style="List Bullet")

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            text = docx_ingestor.extract_text(f.name)

            assert "Apples" in text
            assert "Bananas" in text
            assert "Milk" in text
            os.unlink(f.name)


class TestPdfIngestor:
    def test_can_handle(self):
        assert pdf_ingestor.can_handle("paper.pdf") == True
        assert pdf_ingestor.can_handle("paper.PDF") == True
        assert pdf_ingestor.can_handle("paper.docx") == False


class TestRegistry:
    def test_is_supported(self):
        assert is_supported("doc.txt") == True
        assert is_supported("doc.pdf") == True
        assert is_supported("doc.docx") == True
        assert is_supported("doc.csv") == True
        assert is_supported("doc.html") == True
        assert is_supported("code.py") == True
        # not supported
        assert is_supported("image.png") == False
        assert is_supported("music.mp3") == False

    def test_ingest_txt_file(self):
        text = ("This is a paragraph about dogs. " * 30 + "\n\n") * 5
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
            f.write(text)
            f.flush()
            chunks = ingest(f.name, chunk_size=50, overlap=10)
            assert len(chunks) > 1
            assert all("text" in c for c in chunks)
            assert all("src" in c for c in chunks)
            assert all("format" in c for c in chunks)
            assert chunks[0]["format"] == "text"
            os.unlink(f.name)

    def test_ingest_csv_file(self):
        csv_data = "Product,Price,Category\nLaptop,999,Electronics\nBook,15,Education\nPen,2,Stationery\n"
        with tempfile.NamedTemporaryFile(mode="w", suffix=".csv", delete=False) as f:
            f.write(csv_data)
            f.flush()
            chunks = ingest(f.name, chunk_size=100)
            assert len(chunks) >= 1
            # csv should be converted to readable text
            full_text = " ".join(c["text"] for c in chunks)
            assert "Product: Laptop" in full_text
            assert chunks[0]["format"] == "csv"
            os.unlink(f.name)

    def test_ingest_markdown(self):
        md = "# Title\n\nFirst paragraph about something.\n\n## Sub heading\n\nMore text content here.\n"
        with tempfile.NamedTemporaryFile(mode="w", suffix=".md", delete=False) as f:
            f.write(md)
            f.flush()
            chunks = ingest(f.name)
            assert len(chunks) >= 1
            full_text = " ".join(c["text"] for c in chunks)
            assert "Title" in full_text
            assert chunks[0]["format"] == "markdown"
            os.unlink(f.name)

    def test_ingest_docx_file(self):
        """only runs if python-docx is installed."""
        try:
            from docx import Document
        except ImportError:
            return

        doc = Document()
        doc.add_heading("Project Report", level=1)
        for i in range(10):
            doc.add_paragraph(
                f"This is paragraph {i+1} of the test document. "
                f"It contains enough words to make chunking interesting. "
                f"We are testing the full pipeline from docx to chunks. " * 3
            )

        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
            chunks = ingest(f.name, chunk_size=100, overlap=20)

            assert len(chunks) > 1
            assert chunks[0]["format"] == "docx"
            assert "Project Report" in chunks[0]["text"]
            os.unlink(f.name)

    def test_extract_raw_text(self):
        with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
            f.write("just some raw text for testing")
            f.flush()
            raw = extract_raw_text(f.name)
            assert "raw text" in raw
            os.unlink(f.name)

    def test_unsupported_format_raises(self):
        with tempfile.NamedTemporaryFile(suffix=".xyz", delete=False) as f:
            f.write(b"binary gibberish")
            f.flush()
            try:
                ingest(f.name)
                assert False, "should have raised"
            except ValueError as e:
                assert "dont know" in str(e).lower() or "unsupported" in str(e).lower()
            os.unlink(f.name)
