"""
test_locally.py - run this on your machine to verify everything works.

setup:
    cd tgvectordb/
    pip install -e ".[all,dev]"

then:
    python tests/test_locally.py

this creates sample pdf and docx files, runs them through the
ingestors, and shows you exactly what comes out. no telegram
connection needed - this only tests the offline parts.
"""

import os
import sys
import tempfile
import json
from pathlib import Path

# make sure we can import our package
sys.path.insert(0, str(Path(__file__).parent.parent))


def print_header(text):
    print(f"\n{'='*60}")
    print(f"  {text}")
    print(f"{'='*60}\n")


def print_chunks(chunks, max_text=120):
    for i, chunk in enumerate(chunks):
        txt = chunk["text"][:max_text]
        if len(chunk["text"]) > max_text:
            txt += "..."
        print(f"  chunk {i}: [{chunk.get('format', '?')}] ({len(chunk['text'])} chars)")
        print(f"    src: {chunk.get('src', '?')}")
        print(f"    text: {txt}")
        print()


# ── TEST 1: plain text ──────────────────────────────────

def test_txt():
    print_header("TEST: Plain Text (.txt)")

    from tgvectordb.ingestors.registry import ingest

    text = """Introduction to Machine Learning

Machine learning is a subset of artificial intelligence that focuses on
building systems that learn from data. Instead of being explicitly
programmed to perform a task, these systems use algorithms to parse data,
learn from it, and then make predictions or decisions.

Types of Machine Learning

There are three main types of machine learning: supervised learning,
unsupervised learning, and reinforcement learning. Supervised learning
uses labeled training data to learn the mapping between inputs and outputs.
Unsupervised learning finds hidden patterns in data without labels.
Reinforcement learning trains agents to make sequences of decisions.

Applications

Machine learning is used in many real-world applications including
image recognition, natural language processing, recommendation systems,
fraud detection, autonomous vehicles, and medical diagnosis. The field
continues to grow rapidly with new breakthroughs happening regularly.
"""

    with tempfile.NamedTemporaryFile(mode="w", suffix=".txt", delete=False) as f:
        f.write(text)
        f.flush()
        chunks = ingest(f.name, chunk_size=50, overlap=10)
        print(f"  file: {f.name}")
        print(f"  chunks produced: {len(chunks)}")
        print()
        print_chunks(chunks)
        os.unlink(f.name)

    print("  ✓ text ingestor works!")
    return True


# ── TEST 2: CSV ─────────────────────────────────────────

def test_csv():
    print_header("TEST: CSV (.csv)")

    from tgvectordb.ingestors.registry import ingest

    csv_data = """Name,Department,Role,Experience
Rahul Sharma,Engineering,Backend Developer,5 years
Priya Patel,Design,UI/UX Designer,3 years
Amit Singh,Engineering,Frontend Developer,4 years
Sneha Gupta,Marketing,Content Manager,2 years
Vikram Reddy,Engineering,DevOps Engineer,6 years
"""

    with tempfile.NamedTemporaryFile(mode="w", suffix=".csv", delete=False) as f:
        f.write(csv_data)
        f.flush()
        chunks = ingest(f.name, chunk_size=100)
        print(f"  chunks produced: {len(chunks)}")
        print()
        print_chunks(chunks)
        os.unlink(f.name)

    print("  ✓ csv ingestor works!")
    return True


# ── TEST 3: HTML ────────────────────────────────────────

def test_html():
    print_header("TEST: HTML (.html)")

    from tgvectordb.ingestors.registry import ingest

    html = """<!DOCTYPE html>
<html>
<head><title>Test Page</title>
<style>body { font-family: sans-serif; }</style>
<script>console.log("this should be removed");</script>
</head>
<body>
<h1>Welcome to TgVectorDB</h1>
<p>This is a <strong>test page</strong> with some HTML content.</p>
<h2>Features</h2>
<ul>
<li>Free unlimited storage</li>
<li>Semantic search</li>
<li>No server required</li>
</ul>
<p>Visit our <a href="https://example.com">website</a> for more info.</p>
<div class="footer">Copyright 2026</div>
</body>
</html>"""

    with tempfile.NamedTemporaryFile(mode="w", suffix=".html", delete=False) as f:
        f.write(html)
        f.flush()
        chunks = ingest(f.name, chunk_size=100)
        print(f"  chunks produced: {len(chunks)}")
        print()
        print_chunks(chunks)

        # verify tags are stripped
        full = " ".join(c["text"] for c in chunks)
        assert "<h1>" not in full, "html tags not stripped!"
        assert "console.log" not in full, "script content not removed!"
        os.unlink(f.name)

    print("  ✓ html ingestor works!")
    return True


# ── TEST 4: DOCX ───────────────────────────────────────

def test_docx():
    print_header("TEST: Word Document (.docx)")

    try:
        from docx import Document
    except ImportError:
        print("  ⚠ python-docx not installed. install with:")
        print("    pip install python-docx")
        print("  skipping docx test.")
        return False

    from tgvectordb.ingestors.registry import ingest

    # create a real docx with various elements
    doc = Document()

    doc.add_heading("Quarterly Business Report", level=1)
    doc.add_paragraph(
        "This report summarizes the key metrics and achievements "
        "for Q3 2025. Overall performance exceeded expectations "
        "with significant growth across all departments."
    )

    doc.add_heading("Revenue Summary", level=2)
    doc.add_paragraph(
        "Total revenue for Q3 reached 4.2 million dollars, representing "
        "a 15% increase over the previous quarter. The growth was primarily "
        "driven by new enterprise contracts and expansion of existing accounts."
    )

    # add a table
    doc.add_heading("Department Performance", level=2)
    table = doc.add_table(rows=4, cols=3)
    table.style = "Table Grid"
    headers = ["Department", "Revenue", "Growth"]
    for i, h in enumerate(headers):
        table.cell(0, i).text = h
    data = [
        ["Engineering", "$1.8M", "+18%"],
        ["Sales", "$1.5M", "+12%"],
        ["Marketing", "$0.9M", "+22%"],
    ]
    for row_idx, row_data in enumerate(data):
        for col_idx, val in enumerate(row_data):
            table.cell(row_idx + 1, col_idx).text = val

    doc.add_heading("Key Highlights", level=2)
    doc.add_paragraph("Launched new product line in Southeast Asian markets", style="List Bullet")
    doc.add_paragraph("Hired 45 new engineers across three offices", style="List Bullet")
    doc.add_paragraph("Customer satisfaction score reached all-time high of 94%", style="List Bullet")
    doc.add_paragraph("Successfully migrated infrastructure to cloud", style="List Bullet")

    doc.add_heading("Outlook", level=2)
    doc.add_paragraph(
        "Looking ahead to Q4, we expect continued momentum. The pipeline "
        "is strong with several large deals expected to close before year end. "
        "Key initiatives include expanding our AI capabilities and entering "
        "two new geographic markets."
    )

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc.save(f.name)
        print(f"  created test docx: {f.name}")
        print(f"  file size: {os.path.getsize(f.name)} bytes")
        print()

        # test raw extraction
        from tgvectordb.ingestors.docx_ingestor import extract_text
        raw = extract_text(f.name)
        print(f"  raw extracted text ({len(raw)} chars):")
        print(f"  ---")
        for line in raw.split("\n")[:15]:
            if line.strip():
                print(f"    {line[:100]}")
        print(f"  ---")
        print()

        # test full ingestion pipeline
        chunks = ingest(f.name, chunk_size=80, overlap=15)
        print(f"  chunks produced: {len(chunks)}")
        print()
        print_chunks(chunks)

        # verify key content made it through
        full = " ".join(c["text"] for c in chunks)
        assert "Quarterly" in full, "heading missing"
        assert "revenue" in full.lower(), "body text missing"
        assert "Engineering" in full, "table content missing"
        assert chunks[0]["format"] == "docx"

        os.unlink(f.name)

    print("  ✓ docx ingestor works!")
    return True


# ── TEST 5: PDF ─────────────────────────────────────────

def test_pdf():
    print_header("TEST: PDF (.pdf)")

    try:
        import pdfplumber
    except ImportError:
        print("  ⚠ pdfplumber not installed. install with:")
        print("    pip install pdfplumber")
        print("  skipping pdf test.")
        return False

    # we need reportlab to CREATE a test pdf
    try:
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet
        from reportlab.lib import colors
    except ImportError:
        print("  ⚠ reportlab not installed (needed to create test pdf). install with:")
        print("    pip install reportlab")
        print("  skipping pdf test.")
        return False

    from tgvectordb.ingestors.registry import ingest

    # create a real multi-page pdf
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
        pdf_path = f.name

    doc = SimpleDocTemplate(pdf_path, pagesize=letter)
    styles = getSampleStyleSheet()
    story = []

    story.append(Paragraph("Research Paper: Vector Databases in Modern AI", styles["Title"]))
    story.append(Spacer(1, 20))

    story.append(Paragraph("Abstract", styles["Heading2"]))
    story.append(Paragraph(
        "This paper explores the role of vector databases in modern artificial "
        "intelligence applications. We examine how embedding-based storage and "
        "retrieval systems have become essential for semantic search, "
        "recommendation engines, and retrieval-augmented generation pipelines. "
        "Our analysis covers architectural patterns, performance trade-offs, "
        "and cost considerations for deploying vector databases at scale.",
        styles["Normal"]
    ))
    story.append(Spacer(1, 12))

    story.append(Paragraph("1. Introduction", styles["Heading2"]))
    story.append(Paragraph(
        "The rapid growth of large language models has created an unprecedented "
        "demand for efficient vector storage solutions. Traditional relational "
        "databases struggle with high-dimensional similarity search, leading to "
        "the emergence of specialized vector databases. These systems store data "
        "as numerical embeddings and support approximate nearest neighbor search "
        "with sub-millisecond latency at scale.",
        styles["Normal"]
    ))
    story.append(Spacer(1, 12))

    story.append(Paragraph("2. Methodology", styles["Heading2"]))
    story.append(Paragraph(
        "We evaluated seven popular vector database solutions across three "
        "dimensions: query latency, storage efficiency, and operational complexity. "
        "Each system was tested with datasets ranging from 100,000 to 10 million "
        "vectors at 384 and 1536 dimensions. Benchmarks were conducted on "
        "standardized hardware to ensure fair comparison.",
        styles["Normal"]
    ))
    story.append(Spacer(1, 12))

    # add a table
    story.append(Paragraph("Table 1: Benchmark Results", styles["Heading3"]))
    table_data = [
        ["Database", "Latency (p50)", "Storage/1M vec", "Free Tier"],
        ["Pinecone", "8ms", "6.2 GB", "2 GB"],
        ["Qdrant", "5ms", "4.8 GB", "1 GB"],
        ["ChromaDB", "12ms", "5.1 GB", "Self-hosted"],
        ["TgVectorDB", "500ms cold / 5ms warm", "0 GB local", "Unlimited"],
    ]
    t = Table(table_data)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.grey),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.whitesmoke),
        ("GRID", (0, 0), (-1, -1), 1, colors.black),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
    ]))
    story.append(t)
    story.append(Spacer(1, 12))

    story.append(Paragraph("3. Conclusion", styles["Heading2"]))
    story.append(Paragraph(
        "Our findings suggest that the optimal choice of vector database depends "
        "heavily on the specific use case. For personal and prototyping use cases, "
        "cost-free solutions provide sufficient performance. For production "
        "deployments serving thousands of concurrent users, managed services "
        "with guaranteed SLAs remain the preferred choice.",
        styles["Normal"]
    ))

    doc.build(story)

    print(f"  created test pdf: {pdf_path}")
    print(f"  file size: {os.path.getsize(pdf_path)} bytes")
    print()

    # test raw extraction
    from tgvectordb.ingestors.pdf_ingestor import extract_text
    raw = extract_text(pdf_path)
    print(f"  raw extracted text ({len(raw)} chars):")
    print(f"  ---")
    for line in raw.split("\n")[:12]:
        if line.strip():
            print(f"    {line[:100]}")
    print(f"  ---")
    print()

    # test full ingestion
    chunks = ingest(pdf_path, chunk_size=80, overlap=15)
    print(f"  chunks produced: {len(chunks)}")
    print()
    print_chunks(chunks)

    # verify
    full = " ".join(c["text"] for c in chunks)
    assert "vector" in full.lower(), "main content missing"
    assert chunks[0]["format"] == "pdf"

    os.unlink(pdf_path)

    print("  ✓ pdf ingestor works!")
    return True


# ── TEST 6: quantizer + serialization roundtrip ─────────

def test_quantizer_roundtrip():
    print_header("TEST: Quantizer + Serialization Roundtrip")

    import numpy as np
    from tgvectordb.embedding.quantizer import Quantizer
    from tgvectordb.utils.serialization import pack_vector_message, unpack_vector_message

    q = Quantizer(dimensions=384)

    # simulate what happens when you add a text
    fake_vector = np.random.randn(384).astype(np.float32)
    int8_vec, params = q.quantize(fake_vector)

    metadata = {"src": "test.pdf", "page": 3, "chunk_idx": 0}
    text = "this is a test chunk about vector databases and telegram storage"

    # pack into message format
    msg_str = pack_vector_message(int8_vec, params, metadata, text=text)
    print(f"  message length: {len(msg_str)} / 4096 chars")
    print(f"  fits in telegram: {'yes' if len(msg_str) <= 4096 else 'NO!'}")

    # unpack
    result = unpack_vector_message(msg_str)
    restored_vec = q.dequantize(result["vector_int8"], result["quant_params"])

    # check quality
    cos_sim = np.dot(fake_vector, restored_vec) / (
        np.linalg.norm(fake_vector) * np.linalg.norm(restored_vec)
    )
    print(f"  quantization quality (cosine sim): {cos_sim:.6f}")
    print(f"  metadata preserved: {result['metadata']['src'] == 'test.pdf'}")
    print(f"  text preserved: {result['metadata']['text'][:50]}...")

    assert cos_sim > 0.95, f"quality too low: {cos_sim}"
    assert len(msg_str) <= 4096

    print("\n  ✓ quantizer + serialization works!")
    return True


# ── TEST 7: embedding model (optional, downloads model) ──

def test_embedding():
    print_header("TEST: Embedding Model (e5-small-v2)")
    print("  this will download the model (~130MB) on first run.")
    print("  skip? press Ctrl+C within 5 seconds...")

    import time
    try:
        for i in range(5, 0, -1):
            print(f"  starting in {i}...", end="\r")
            time.sleep(1)
        print("  starting now...        ")
    except KeyboardInterrupt:
        print("\n  skipped.")
        return False

    from tgvectordb.embedding.model import EmbeddingModel
    import numpy as np

    model = EmbeddingModel()

    # embed a query
    qvec = model.embed_query("how do plants make food?")
    print(f"  query vector shape: {qvec.shape}")
    print(f"  query vector dtype: {qvec.dtype}")
    print(f"  query vector sample: [{qvec[0]:.4f}, {qvec[1]:.4f}, {qvec[2]:.4f}, ...]")

    # embed some documents
    docs = [
        "Photosynthesis converts sunlight into energy in plants",
        "Machine learning uses neural networks for pattern recognition",
        "The capital of France is Paris",
    ]
    doc_vecs = model.embed_documents_batch(docs)
    print(f"  doc vectors shape: {doc_vecs.shape}")

    # check that similar things are actually similar
    from tgvectordb.search.engine import cosine_similarity_batch
    sims = cosine_similarity_batch(qvec, doc_vecs)
    print(f"\n  similarity scores for 'how do plants make food?':")
    for i, (doc, sim) in enumerate(zip(docs, sims)):
        marker = " ← best match!" if sim == max(sims) else ""
        print(f"    [{sim:.4f}] {doc[:60]}{marker}")

    # the photosynthesis doc should be the best match
    assert np.argmax(sims) == 0, "wrong best match!"

    print("\n  ✓ embedding model works correctly!")
    return True


# ── RUN EVERYTHING ──────────────────────────────────────

if __name__ == "__main__":
    print("\n" + "=" * 60)
    print("  TgVectorDB Local Test Suite")
    print("  tests ingestors, quantizer, and embedding (no telegram needed)")
    print("=" * 60)

    results = {}

    tests = [
        ("Plain Text", test_txt),
        ("CSV", test_csv),
        ("HTML", test_html),
        ("DOCX", test_docx),
        ("PDF", test_pdf),
        ("Quantizer", test_quantizer_roundtrip),
        ("Embedding", test_embedding),
    ]

    for name, test_fn in tests:
        try:
            passed = test_fn()
            results[name] = "✓ passed" if passed else "⚠ skipped (missing dependency)"
        except Exception as e:
            results[name] = f"✗ FAILED: {e}"
            import traceback
            traceback.print_exc()

    # summary
    print_header("RESULTS")
    for name, status in results.items():
        print(f"  {name:20s} {status}")

    failed = sum(1 for v in results.values() if v.startswith("✗"))
    if failed:
        print(f"\n  {failed} test(s) failed!")
        sys.exit(1)
    else:
        print(f"\n  all tests passed! ready to connect to telegram.")
        print(f"  next step: run examples/quickstart.py with your credentials.")
